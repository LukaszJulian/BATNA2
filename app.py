import streamlit as st
import anthropic
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Page configuration
st.set_page_config(
    page_title="BATNA Document Creator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
def initialize_session_state():
    if 'collected_data' not in st.session_state:
        st.session_state.collected_data = {}
    if 'final_document' not in st.session_state:
        st.session_state.final_document = None
    if 'show_document' not in st.session_state:
        st.session_state.show_document = False
    if 'document_history' not in st.session_state:
        st.session_state.document_history = []
    # Initialize form inputs
    sections = [
        "negotiation_subject", "project_value", "company_profile", 
        "scope_description", "targets", "vendors", "interests", 
        "advantages", "disadvantages"
    ]
    for section in sections:
        if f"input_{section}" not in st.session_state:
            st.session_state[f"input_{section}"] = ""

# Initialize the Anthropic client
def init_client():
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
        if not api_key:
            st.error("API key not found in secrets.")
            st.stop()
            return None
            
        client = anthropic.Anthropic(api_key=api_key)
        return client
    except KeyError:
        st.error("'ANTHROPIC_API_KEY' not found in Streamlit secrets. Please check your secrets configuration.")
        st.stop()
        return None
    except Exception as e:
        st.error(f"Failed to initialize Anthropic client. Error: {str(e)}")
        st.stop()
        return None

def get_assistant_response(client, prompt):
    try:
        response = client.messages.create(
            model="claude-3-opus-20240229",
            messages=[{
                "role": "user",
                "content": prompt
            }],
            max_tokens=4096,
            temperature=0.7
        )
        return response.content[0].text
    except Exception as e:
        st.error(f"Error getting response from Claude 3 Opus: {str(e)}")
        st.write("Detailed error:", str(e))
        return None

def save_to_history(document, data):
    """Save generated document to history with timestamp"""
    # Keep only the 10 most recent documents
    if len(st.session_state.document_history) >= 10:
        st.session_state.document_history.pop(0)  # Remove oldest document
        
    history_entry = {
        'timestamp': datetime.now(),
        'document': document,
        'metadata': {
            'subject': data.get('negotiation_subject', 'Untitled'),
            'value': data.get('project_value', 'N/A'),
            'data': data.copy()
        }
    }
    st.session_state.document_history.append(history_entry)

def render_sidebar_history():
    with st.sidebar:
        st.title("Recent Documents")
        
        if not st.session_state.document_history:
            st.info("No documents generated yet")
            return

        # Display documents from newest to oldest
        for idx, entry in enumerate(reversed(st.session_state.document_history)):
            with st.expander(
                f"📄 {entry['metadata']['subject'][:40]}... \n"
                f"{entry['timestamp'].strftime('%Y-%m-%d %H:%M')}"
            ):
                st.write(f"**Project Value:** {entry['metadata']['value']}")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("View", key=f"view_{idx}", use_container_width=True):
                        st.session_state.final_document = entry['document']
                        st.session_state.collected_data = entry['metadata']['data']
                        st.session_state.show_document = True
                        st.rerun()
                
                with col2:
                    if st.button("Delete", key=f"delete_{idx}", use_container_width=True):
                        st.session_state.document_history.remove(entry)
                        st.rerun()
        
        if st.session_state.document_history:
            if st.button("Clear History", use_container_width=True):
                st.session_state.document_history = []
                st.rerun()

# Document Generation Functions
def generate_pdf(content, filename):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Add custom styles
    styles.add(ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=16
    ))
    
    styles.add(ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=12
    ))
    
    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        spaceAfter=30
    )
    story.append(Paragraph("BATNA Analysis Document", title_style))
    story.append(Spacer(1, 12))
    
    try:
        # Split content into sections
        sections = content.split('\n')
        current_text = ""
        
        for line in sections:
            line = line.strip()
            if not line:
                if current_text:
                    # Clean the text
                    cleaned_text = (current_text
                        .replace('|', ' - ')
                        .replace('<br>', '\n')
                        .replace('>', '')
                        .replace('<', '')
                    )
                    
                    # Check if it's a heading
                    if cleaned_text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
                        story.append(Paragraph(cleaned_text, styles['CustomHeading']))
                    else:
                        story.append(Paragraph(cleaned_text, styles['CustomBody']))
                    
                    story.append(Spacer(1, 8))
                    current_text = ""
            else:
                current_text += " " + line if current_text else line
        
        # Add any remaining text
        if current_text:
            cleaned_text = (current_text
                .replace('|', ' - ')
                .replace('<br>', '\n')
                .replace('>', '')
                .replace('<', '')
            )
            story.append(Paragraph(cleaned_text, styles['CustomBody']))
    
    except Exception as e:
        st.error(f"Error processing content for PDF: {str(e)}")
        return None
    
    try:
        doc.build(story)
        pdf = buffer.getvalue()
        buffer.close()
        return pdf
    except Exception as e:
        st.error(f"Error building PDF: {str(e)}")
        return None

def generate_word_doc(content, filename):
    """Generate a formatted Word document"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('BATNA Analysis Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(datetime.now().strftime("%Y-%m-%d")).italic = True
    
    doc.add_paragraph()  # Add spacing
    
    try:
        # Split content into sections
        sections = content.split('\n')
        current_text = ""
        
        for line in sections:
            line = line.strip()
            if not line:
                if current_text:
                    # Clean the text
                    cleaned_text = (current_text
                        .replace('|', ' - ')
                        .replace('<br>', '\n')
                        .replace('>', '')
                        .replace('<', '')
                    )
                    
                    # Check if it's a heading
                    if cleaned_text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
                        doc.add_heading(cleaned_text, level=1)
                    else:
                        doc.add_paragraph(cleaned_text)
                    
                    current_text = ""
            else:
                current_text += " " + line if current_text else line
        
        # Add any remaining text
        if current_text:
            cleaned_text = (current_text
                .replace('|', ' - ')
                .replace('<br>', '\n')
                .replace('>', '')
                .replace('<', '')
            )
            doc.add_paragraph(cleaned_text)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error generating Word document: {str(e)}")
        return None

def generate_text_doc(content):
    """Generate a plain text document"""
    try:
        # Clean and format the content
        cleaned_content = (content
            .replace('|', ' - ')
            .replace('<br>', '\n')
            .replace('>', '')
            .replace('<', '')
        )
        
        # Add header
        text_content = f"""BATNA ANALYSIS DOCUMENT
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
{'='*50}\n\n"""
        
        text_content += cleaned_content
        
        return text_content.encode('utf-8')
        
    except Exception as e:
        st.error(f"Error generating text document: {str(e)}")
        return None

def render_input_form():
    sections = {
        "negotiation_subject": "Negotiation Subject",
        "project_value": "Project Value",
        "company_profile": "Company's Profile & Industry",
        "scope_description": "Scope Description",
        "targets": "Targets to be Achieved",
        "vendors": "Vendors & Suppliers to be Invited",
        "interests": "Client's & Vendor's Interests",
        "advantages": "Client's & Vendors' Negotiation Advantages",
        "disadvantages": "Client's & Vendors' Negotiation Disadvantages"
    }
    
    # Create form for all inputs
    with st.form("batna_form"):
        # Create two columns for inputs
        col1, col2 = st.columns(2)
        
        # Split sections between columns
        sections_list = list(sections.items())
        mid_point = len(sections_list) // 2 + len(sections_list) % 2
        
        all_fields_filled = True
        
        # First column
        with col1:
            for key, title in sections_list[:mid_point]:
                st.subheader(title)
                value = st.text_area(
                    "Enter information:",
                    value=st.session_state.get(f"input_{key}", ""),
                    height=150,
                    key=f"input_{key}",
                    label_visibility="collapsed"
                )
                
                if not value.strip():
                    all_fields_filled = False
                st.markdown("---")
        
        # Second column
        with col2:
            for key, title in sections_list[mid_point:]:
                st.subheader(title)
                value = st.text_area(
                    "Enter information:",
                    value=st.session_state.get(f"input_{key}", ""),
                    height=150,
                    key=f"input_{key}",
                    label_visibility="collapsed"
                )
                
                if not value.strip():
                    all_fields_filled = False
                st.markdown("---")
        
        # Submit button
        submitted = st.form_submit_button("Generate BATNA Document", use_container_width=True)
        if submitted:
            if all_fields_filled:
                for key in sections.keys():
                    st.session_state.collected_data[key] = st.session_state[f"input_{key}"]
                generate_batna_document()
            else:
                st.error("Please fill in all fields before generating the BATNA document.")

def generate_batna_document():
    client = init_client()
    if not client:
        return

    prompt = """Create a comprehensive BATNA document based on the following information:

    {input_data}

    As an AI assistant, your task is to provide a comprehensive, well-structured analysis of a negotiation scenario between a client and a vendor. The analysis should cover various aspects, including the client's and vendor's Best Alternative To a Negotiated Agreement (BATNA), risk assessment and mitigation strategies, negotiation strategy and tactics, and an implementation roadmap. The analysis should be thorough, considering multiple scenarios and potential outcomes, and should include actionable recommendations for the client.
    Please use for the final BATNA Document the structure below:

    1. EXECUTIVE SUMMARY
    [In this section, provide a concise yet comprehensive overview of the entire analysis, highlighting the key findings and critical recommendations. The summary should be engaging and persuasive, encouraging the reader to delve into the details of the analysis.]

    2. CLIENT'S BATNA ANALYSIS
    [Format: table with following columns: Alternative Options, Strength Assessment, Risk Assessment. Instruction: Following the table, present a detailed analysis of each viable alternative available to the client. Consider factors such as implementation feasibility, cost implications, and strategic fit. The analysis should provide a clear understanding of the client's position and the potential outcomes of pursuing alternative options.]

    3. VENDOR'S BATNA ANALYSIS
    [Format: table with following columns: Alternative Options, Strength Assessment, Risk Assessment. Instruction: After the table, analyze the vendor's alternatives, market position, and potential responses to different scenarios. This analysis should provide insights into the vendor's strengths, weaknesses, and likely negotiation strategies.]

    4. RISK ASSESSMENT & MITIGATION
    [Format: table with following columns: Risk Category, Mitigation Strategy, Contingency Plan. Instruction: Following the table, provide a comprehensive analysis of potential risks associated with the negotiation and the proposed mitigation strategies. Consider various risk categories, such as financial, operational, and reputational risks. The analysis should demonstrate a proactive approach to risk management and contingency planning.]

    5. NEGOTIATION STRATEGY
    [Present a detailed strategic approach to the negotiation, covering core objectives, non-negotiables, value creation opportunities, power dynamics, leverage points, and relationship management. The strategy should be well-reasoned and adaptable, taking into account the insights gained from the BATNA and risk analyses.]

    6. NEGOTIATION TACTICS
    [Detail specific tactical recommendations for the negotiation, including type of the process (e.g. single source, RFx, auction...), opening positions, communication strategies, response scenarios, and timing considerations. The tactics should be aligned with the overall negotiation strategy and designed to maximize the client's chances of achieving a favorable outcome.]

    7. RECOMMENDATIONS
    [Provide clear, actionable recommendations for the client, including immediate next steps, critical success factors, resource requirements, and expected outcomes. The recommendations should be based on the insights gained from the analysis and designed to help the client achieve their objectives in the negotiation.]
    
    Formating:  
    -heading"bold and larger than normal text
    -devider after every of the 8 sections
    """
    
    formatted_data = "\n\n".join([f"{k.replace('_', ' ').title()}: {v}" 
                                 for k, v in st.session_state.collected_data.items()])
    
    with st.spinner("Generating BATNA document..."):
        response = get_assistant_response(client, prompt.format(input_data=formatted_data))
        if response:
            st.session_state.final_document = response
            save_to_history(response, st.session_state.collected_data)
            st.session_state.show_document = True
            st.rerun()

def main():
    st.title("BATNA Assistant")
    
    # Professional description
    st.markdown("""
        <div style='background-color: #EBF5FB; padding: 1rem; border-radius: 10px; border-left: 5px solid #2E86C1; margin-bottom: 2rem;'>
            <p style='color: #2E86C1; font-size: 1.1em; margin-bottom: 0.5rem;'><strong>Professional BATNA Document Generator</strong></p>
            <p style='font-size: 0.95em; color: #34495E; line-height: 1.5;'>
                Leverage advanced AI to create detailed Best Alternative To a Negotiated Agreement (BATNA) documents. 
                Our assistant helps procurement professionals analyze negotiation positions, assess risks, 
                and develop strategic alternatives for successful negotiations.
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    initialize_session_state()
    render_sidebar_history()
    
    if not st.session_state.show_document:
        render_input_form()
    else:
        # Display generated document
        st.header("📄 Generated BATNA Document")
        st.markdown("---")
        
        if st.session_state.final_document:
            # Add expander to show input data
            with st.expander("View Input Data", expanded=False):
                for key, value in st.session_state.collected_data.items():
                    st.markdown(f"**{key.replace('_', ' ').title()}:**")
                    st.markdown(f"{value}")
                    st.markdown("---")
            
            st.markdown(st.session_state.final_document)
            
            # Export options
            st.markdown("---")
            st.subheader("Export Options")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("Create New Document", use_container_width=True):
                    st.session_state.collected_data = {}
                    st.session_state.show_document = False
                    st.rerun()
            
            with col2:
                if st.download_button(
                    label="Download as PDF",
                    data=generate_pdf(st.session_state.final_document, "BATNA_document.pdf"),
                    file_name=f"BATNA_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                ):
                    st.success("PDF downloaded successfully!")
            
            with col3:
                if st.download_button(
                    label="Download as Word",
                    data=generate_word_doc(st.session_state.final_document, "BATNA_document.docx"),
                    file_name=f"BATNA_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                ):
                    st.success("Word document downloaded successfully!")
            
            with col4:
                if st.download_button(
                    label="Download as Text",
                    data=generate_text_doc(st.session_state.final_document),
                    file_name=f"BATNA_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                ):
                    st.success("Text document downloaded successfully!")

if __name__ == "__main__":
    main()
